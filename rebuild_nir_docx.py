#!/usr/bin/env python3
"""Собрать Word‑черновик НИР из markdown с кликабельными ссылками."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x4E, 0x79)
MD = Path("НИР_ЭДО_полный_черновик.md")
OUT = Path("НИР_ЭДО_полный_черновик.docx")
LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")
BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def set_run(run, size=14, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_hyperlink(paragraph, text, url, size=14):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_mixed(paragraph, text, size=14, bold=False):
    pos = 0
    for m in LINK_RE.finditer(text):
        if m.start() > pos:
            chunk = text[pos : m.start()]
            add_bold_runs(paragraph, chunk, size, bold)
        add_hyperlink(paragraph, m.group(1), m.group(2), size=size)
        pos = m.end()
    if pos < len(text):
        add_bold_runs(paragraph, text[pos:], size, bold)


def add_bold_runs(paragraph, text, size, bold):
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run(run, size=size, bold=bold)
        run = paragraph.add_run(m.group(1))
        set_run(run, size=size, bold=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run(run, size=size, bold=bold)


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(14)
    styles["Normal"].paragraph_format.line_spacing = 1.5
    for hn, sz in (("Title", 18), ("Heading 1", 16), ("Heading 2", 15), ("Heading 3", 14)):
        st = styles[hn]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = NAVY

    lines = MD.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        i += 1
        if not line.strip() or line.strip() == "---":
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            for run in p.runs:
                set_run(run, size=16, bold=True, color=NAVY)
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            for run in p.runs:
                set_run(run, size=15, bold=True, color=NAVY)
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            for run in p.runs:
                set_run(run, size=14, bold=True, color=NAVY)
            continue
        if line.startswith("|"):
            table_lines = [line]
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            rows = [r.strip("|") for r in table_lines if not re.match(r"^\|?\s*-+", r)]
            parsed = [[c.strip() for c in r.split("|")] for r in rows]
            if parsed:
                table = doc.add_table(rows=len(parsed), cols=len(parsed[0]))
                table.style = "Table Grid"
                for r, row in enumerate(parsed):
                    for c, val in enumerate(row):
                        cell = table.rows[r].cells[c]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_mixed(p, val, size=11, bold=(r == 0))
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_mixed(p, line[2:], size=14)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        add_mixed(p, line, size=14)

    doc.save(OUT)
    print("wrote", OUT, OUT.stat().st_size)


if __name__ == "__main__":
    main()
